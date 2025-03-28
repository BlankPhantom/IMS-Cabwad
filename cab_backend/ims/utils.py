from docx import Document
from django.conf import settings
import os
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

def replace_text_in_table(table, placeholders):
    """Replaces placeholders in a Word table while preserving formatting (bold, italics, etc.)."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Step 1: Rebuild full text from all runs (to detect split placeholders)
                full_text = "".join(run.text for run in paragraph.runs)

                modified_text = full_text  # Copy original text
                replaced = False  # Track if replacement happens

                for key, value in placeholders.items():
                    if key in modified_text:
                        print(f"Replacing {key} with {value}")  # Debugging
                        modified_text = modified_text.replace(key, str(value))
                        replaced = True  # Mark that changes were made
                
                # Step 2: If replacements happened, update runs WITHOUT losing formatting
                if replaced:
                    index = 0  # Track position in modified_text
                    for run in paragraph.runs:
                        run_length = len(run.text)  # Get the original length of this run
                        run.text = modified_text[index: index + run_length]  # Replace only this part
                        index += run_length  # Move to the next part
                        
def generate_reports_doc(report):
    template_path = os.path.join(settings.BASE_DIR, 'Document Format', 'MONTHLY SUMMARY REPORT.docx')
    doc = Document(template_path)

    def format_number(value):
        try:
            value = float(value)
            return f"{value:,.2f}"
        except (TypeError, ValueError):
            return "0.00"

    placeholders = {
        '{{month}}': datetime(1900, int(getattr(report, 'month', 1)+1) if str(getattr(report, 'month', '')).isdigit()+1 else 1, 1).strftime("%B"),
        '{{year}}': str(getattr(report, 'year', datetime.now().year)),
        
        '(tNSC0)': '₱' + format_number(report.totalNSC),
        '(tProd0)': '₱' + format_number(report.totalProd),
        '(tMeterM0)': '₱' + format_number(report.totalMeterMaintenance),
        '(tSpecProj0)': '₱' + format_number(report.totalSpecialProj),
        '(tConst0)': '₱' + format_number(report.totalConstruction),
        '(tComm0)': '₱' + format_number(report.totalCommercial),
        '(tsales0)': '₱' + format_number(report.totalSales),
        '(tGenServe0)': '₱' + format_number(report.totalGenService),
        '(TotalCons0)': '₱' + format_number(report.totalConsumption),

        '(tNSC1)': '₱' + format_number(report.totalNSC1),
        '(tProd1)': '₱' + format_number(report.totalProd1),
        '(tMeterM1)': '₱' + format_number(report.totalMeterMaintenance1),
        '(tSpecProj1)': '₱' + format_number(report.totalSpecialProj1),
        '(tConst1)': '₱' + format_number(report.totalConstruction1),
        '(tComm1)': '₱' + format_number(report.totalCommercial1),
        '(tsales1)': '₱' + format_number(report.totalSales1),
        '(tGenServe1)': '₱' + format_number(report.totalGenService1),
        '(TotalCons1)': '₱' + format_number(report.totalConsumption1),

        '(tNSC2)': '₱' + format_number(report.totalNSC2),
        '(tProd2)': '₱' + format_number(report.totalProd2),
        '(tMeterM2)': '₱' + format_number(report.totalMeterMaintenance2),
        '(tSpecProj2)': '₱' + format_number(report.totalSpecialProj2),
        '(tConst2)': '₱' + format_number(report.totalConstruction2),
        '(tComm2)': '₱' + format_number(report.totalCommercial2),
        '(tsales2)': '₱' + format_number(report.totalSales2),
        '(tGenServe2)': '₱' + format_number(report.totalGenService2),
        '(TotalCons2)': '₱' + format_number(report.totalConsumption2),

        '(tNSC3)': '₱' + format_number(report.totalNSC3),
        '(tProd3)': '₱' + format_number(report.totalProd3),
        '(tMeterM3)': '₱' + format_number(report.totalMeterMaintenance3),
        '(tSpecProj3)': '₱' + format_number(report.totalSpecialProj3),
        '(tConst3)': '₱' + format_number(report.totalConstruction3),
        '(tComm3)': '₱' + format_number(report.totalCommercial3),
        '(tsales3)': '₱' + format_number(report.totalSales3),
        '(tGenServe3)': '₱' + format_number(report.totalGenService3),
        '(TotalCons3)': '₱' + format_number(report.totalConsumption3),

        '(tNSC4)': '₱' + format_number(report.totalNSC4),
        '(tProd4)': '₱' + format_number(report.totalProd4),
        '(tMeterM4)': '₱' + format_number(report.totalMeterMaintenance4),
        '(tSpecProj4)': '₱' + format_number(report.totalSpecialProj4),
        '(tConst4)': '₱' + format_number(report.totalConstruction4),
        '(tComm4)': '₱' + format_number(report.totalCommercial4),
        '(tsales4)': '₱' + format_number(report.totalSales4),
        '(tGenServe4)': '₱' + format_number(report.totalGenService4),
        '(TotalCons4)': '₱' + format_number(report.totalConsumption4),

        '(tNSC5)': '₱' + format_number(report.totalNSC5),
        '(tProd5)': '₱' + format_number(report.totalProd5),
        '(tMeterM5)': '₱' + format_number(report.totalMeterMaintenance5),
        '(tSpecProj5)': '₱' + format_number(report.totalSpecialProj5),
        '(tConst5)': '₱' + format_number(report.totalConstruction5),
        '(tComm5)': '₱' + format_number(report.totalCommercial5),
        '(tsales5)': '₱' + format_number(report.totalSales5),
        '(tGenServe5)': '₱' + format_number(report.totalGenService5),
        '(TotalCons5)': '₱' + format_number(report.totalConsumption5),
    }

    for table in doc.tables:
        replace_text_in_table(table, placeholders)

    output_dir = os.path.join(settings.MEDIA_ROOT, 'generated_reports')
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, f'sample_report_{datetime.now().strftime("%Y%m%d")}.docx')
    
    try:
        doc.save(output_path)
        print(f"Document saved successfully: {output_path}")
    except PermissionError as e:
        print(f"Permission denied: {e}")
        raise
    
    return output_path

def generate_reports_excel(report):
    template_path = os.path.join(settings.BASE_DIR, 'Document Format', 'MONTHLY CONSUMPTION SECTION.xlsx')
    doc = Document(template_path)
    
    placeholders = {
        '{{month}}': datetime.now().strftime("%B"),
        '{{year}}': datetime.now().strftime("%Y"),
    }